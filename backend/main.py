"""
FastAPI application — upload endpoint, job status, and result retrieval.
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import uuid
from pathlib import Path
from urllib.parse import urlparse

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from auth import APIKeyMiddleware
from config import settings
from confluence import ConfluenceClient
from database import JobStatus, create_job, get_db, get_job, init_db, update_job_status
from tasks import process_recording

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MAX_UPLOAD_BYTES = 500 * 1024 * 1024          # 500 MB
CHUNK_SIZE = 256 * 1024                        # 256 KB read chunks
ALLOWED_EXTENSIONS = {".mp4", ".mp3", ".webm", ".wav", ".m4a"}
ALLOWED_CONTENT_TYPES = {
    "video/mp4",
    "audio/mpeg",
    "audio/mp3",
    "video/webm",
    "audio/webm",
    "audio/wav",
    "audio/x-wav",
    "audio/mp4",
    "audio/x-m4a",
}

# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------

app = FastAPI(
    title="Memora API",
    description="Converts meeting recordings into structured Confluence documentation.",
    version="0.1.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.add_middleware(APIKeyMiddleware)


@app.on_event("startup")
async def on_startup() -> None:
    init_db()
    Path(settings.upload_dir).mkdir(parents=True, exist_ok=True)
    logger.info("Memora API ready — uploads dir: %s", settings.upload_dir)


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@app.get("/health", tags=["ops"])
async def health() -> dict:
    """Liveness probe."""
    return {"status": "ok"}


# ---------------------------------------------------------------------------
# Session
# ---------------------------------------------------------------------------

@app.get("/session", tags=["auth"])
async def create_session() -> JSONResponse:
    """
    Issue a short-lived signed session token for browser clients.

    The token is HMAC-signed with ORG_API_KEY and expires after 24 hours.
    The frontend calls this once on load and includes the token via the
    X-Session-Token header — the actual ORG_API_KEY never reaches the browser.
    """
    from auth import _sign_session
    token = _sign_session(settings.org_api_key)
    return JSONResponse(content={"token": token})


# ---------------------------------------------------------------------------
# Confluence destination helpers
# ---------------------------------------------------------------------------

@app.get("/confluence/spaces", tags=["confluence"])
async def list_confluence_spaces() -> JSONResponse:
    """Return available Confluence spaces for the destination picker."""
    try:
        client = ConfluenceClient()
        spaces = client.get_spaces()
        return JSONResponse(content=spaces)
    except Exception as exc:
        logger.error("Failed to fetch Confluence spaces: %s", exc)
        raise HTTPException(status_code=502, detail=f"Could not fetch Confluence spaces: {exc}") from exc


@app.get("/confluence/pages", tags=["confluence"])
async def list_confluence_pages(space_key: str) -> JSONResponse:
    """Return pages in a Confluence space for the parent page picker."""
    if not space_key.strip():
        raise HTTPException(status_code=400, detail="space_key is required.")
    try:
        client = ConfluenceClient()
        pages = client.get_pages(space_key)
        return JSONResponse(content=pages)
    except Exception as exc:
        logger.error("Failed to fetch pages for space %r: %s", space_key, exc)
        raise HTTPException(status_code=502, detail=f"Could not fetch pages: {exc}") from exc


# ---------------------------------------------------------------------------
# POST /upload
# ---------------------------------------------------------------------------

@app.post("/upload", status_code=status.HTTP_202_ACCEPTED, tags=["jobs"])
async def upload_meeting(
    file: UploadFile = File(...),
    output_type: str = Form("detailed"),
    publish_to_confluence: bool = Form(True),
    custom_instructions: str = Form(""),
    confluence_space_key: str = Form(""),
    confluence_parent_page_id: str = Form(""),
    confluence_page_title: str = Form(""),
    context_text: str = Form(""),
    confluence_reference_url: str = Form(""),
    screenshots_enabled: bool = Form(False),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Accept a meeting recording, persist it, create a job record, and queue processing.

    - Allowed formats: .mp4 .mp3 .webm .wav .m4a
    - Maximum size: 500 MB (enforced by streaming — Content-Length is not trusted)
    - Returns immediately with {job_id, status} — poll /status/{job_id} for progress
    """
    original_name = file.filename or "upload"
    ext = Path(original_name).suffix.lower()

    # --- Extension check (first line of defence; MIME checked below) --------
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    # --- MIME type check (second line of defence) ---------------------------
    content_type = (file.content_type or "").split(";")[0].strip().lower()
    if content_type and content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported content type '{content_type}'.",
        )

    # --- Stream to disk, enforcing the 500 MB cap ---------------------------
    file_id = str(uuid.uuid4())
    storage_filename = f"{file_id}{ext}"           # e.g. "a1b2c3d4....mp4"
    dest_path = Path(settings.upload_dir) / storage_filename

    bytes_written = 0
    try:
        with dest_path.open("wb") as out:
            while True:
                chunk = await file.read(CHUNK_SIZE)
                if not chunk:
                    break
                bytes_written += len(chunk)
                if bytes_written > MAX_UPLOAD_BYTES:
                    raise HTTPException(
                        status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                        detail=f"File exceeds the 500 MB limit.",
                    )
                out.write(chunk)
    except HTTPException:
        dest_path.unlink(missing_ok=True)   # clean up partial file
        raise
    except Exception as exc:
        dest_path.unlink(missing_ok=True)
        logger.exception("Unexpected error while saving upload")
        raise HTTPException(status_code=500, detail="Failed to save file.") from exc

    if bytes_written == 0:
        dest_path.unlink(missing_ok=True)
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty.")

    # --- Create DB record ---------------------------------------------------
    job = create_job(
        db,
        filename=original_name,
        storage_path=storage_filename,
        output_type=output_type,
        publish_to_confluence=publish_to_confluence,
        custom_instructions=custom_instructions.strip() or None,
        confluence_space_key=confluence_space_key.strip() or None,
        confluence_parent_page_id=confluence_parent_page_id.strip() or None,
        confluence_page_title=confluence_page_title.strip() or None,
        context_text=context_text.strip() or None,
        confluence_reference_url=confluence_reference_url.strip() or None,
        screenshots_enabled=screenshots_enabled,
    )

    # --- Queue Celery task --------------------------------------------------
    process_recording.delay(job.id)

    logger.info("Job %s queued — file: %s (%d bytes)", job.id, storage_filename, bytes_written)
    return JSONResponse(
        status_code=status.HTTP_202_ACCEPTED,
        content={"job_id": job.id, "status": job.status.value},
    )


# ---------------------------------------------------------------------------
# POST /upload-url
# ---------------------------------------------------------------------------

class UrlSubmitRequest(BaseModel):
    url: str
    title: str = ""
    output_type: str = "detailed"
    publish_to_confluence: bool = True
    custom_instructions: str = ""
    confluence_space_key: str = ""
    confluence_parent_page_id: str = ""
    confluence_page_title: str = ""
    context_text: str = ""
    confluence_reference_url: str = ""
    screenshots_enabled: bool = False


@app.post("/upload-url", status_code=status.HTTP_202_ACCEPTED, tags=["jobs"])
async def upload_from_url(
    body: UrlSubmitRequest,
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Submit a publicly accessible recording URL for processing.

    The file is downloaded inside the Celery worker (not here), so this
    returns immediately.  Poll /status/{job_id} for progress.

    Supported sources: any URL yt-dlp can handle — direct file links,
    Zoom cloud recordings, Vimeo, YouTube, etc.
    """
    parsed = urlparse(body.url)
    if parsed.scheme not in ("http", "https"):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="URL must use http or https.",
        )

    # Derive a human-readable display name from the URL path.
    url_path = parsed.path.rstrip("/")
    display_name = Path(url_path).name or "recording"

    # If the user supplied a title, use that as the display filename.
    if body.title.strip():
        display_name = body.title.strip()

    job = create_job(
        db,
        filename=display_name,
        source_url=body.url,
        output_type=body.output_type,
        publish_to_confluence=body.publish_to_confluence,
        custom_instructions=body.custom_instructions.strip() or None,
        confluence_space_key=body.confluence_space_key.strip() or None,
        confluence_parent_page_id=body.confluence_parent_page_id.strip() or None,
        confluence_page_title=body.confluence_page_title.strip() or None,
        context_text=body.context_text.strip() or None,
        confluence_reference_url=body.confluence_reference_url.strip() or None,
        screenshots_enabled=body.screenshots_enabled,
    )
    process_recording.delay(job.id)

    logger.info("Job %s queued from URL: %s", job.id, body.url)
    return JSONResponse(
        status_code=status.HTTP_202_ACCEPTED,
        content={"job_id": job.id, "status": job.status.value},
    )


# ---------------------------------------------------------------------------
# GET /status/{job_id}
# ---------------------------------------------------------------------------

@app.get("/status/{job_id}", tags=["jobs"])
async def job_status(job_id: str, db: Session = Depends(get_db)) -> JSONResponse:
    """
    Return the current processing status of a job.

    Poll this until status is 'done' or 'failed'. Response shape:
    {
        "job_id": "...",
        "status": "transcribing",          # one of the JobStatus values
        "filename": "meeting.mp4",
        "created_at": "2026-05-16T...",
        "updated_at": "2026-05-16T...",
        "error_message": null              # non-null only when status == "failed"
    }
    """
    try:
        job = get_job(db, job_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="Job not found.")

    agent_decisions: list = []
    if job.result_json:
        try:
            agent_decisions = json.loads(job.result_json).get("agent_decisions") or []
        except json.JSONDecodeError:
            pass

    return JSONResponse(content={
        "job_id":          job.id,
        "status":          job.status.value,
        "filename":        job.filename,
        "created_at":      job.created_at.isoformat() if job.created_at else None,
        "updated_at":      job.updated_at.isoformat() if job.updated_at else None,
        "error_message":   job.error_message,
        "agent_decisions": agent_decisions,
    })


# ---------------------------------------------------------------------------
# GET /result/{job_id}
# ---------------------------------------------------------------------------

@app.get("/result/{job_id}", tags=["jobs"])
async def job_result(job_id: str, db: Session = Depends(get_db)) -> JSONResponse:
    """
    Return the full result of a completed job.

    Returns 404 if the job does not exist.
    Returns 409 if the job has not reached 'done' yet (including 'failed').
    Response shape when done:
    {
        "job_id": "...",
        "status": "done",
        "confluence_url": "https://...",
        "result": {                        # parsed from result_json
            "title": "...",
            "summary": "...",
            "action_items": [...],
            "decisions": [...],
            "attendees": [...]
        }
    }
    """
    try:
        job = get_job(db, job_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="Job not found.")

    if job.status == JobStatus.FAILED:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={"message": "Job failed.", "error": job.error_message},
        )

    if job.status != JobStatus.DONE:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={"message": "Job is not complete yet.", "status": job.status.value},
        )

    try:
        result_data = json.loads(job.result_json) if job.result_json else None
    except json.JSONDecodeError:
        logger.error("Corrupt result_json for job %s", job_id)
        result_data = None

    return JSONResponse(content={
        "job_id":         job.id,
        "status":         job.status.value,
        "confluence_url": job.confluence_url,
        "result":         result_data,
        "publish_failed": bool(job.publish_failed),
    })


# ---------------------------------------------------------------------------
# POST /jobs/{job_id}/retry-publish
# ---------------------------------------------------------------------------

@app.post("/jobs/{job_id}/retry-publish", status_code=status.HTTP_202_ACCEPTED, tags=["jobs"])
async def retry_publish_endpoint(job_id: str, db: Session = Depends(get_db)) -> JSONResponse:
    """
    Re-queue Confluence publishing for a completed job where publish_failed=True.

    Sets status → PUBLISHING and clears publish_failed before queuing the task,
    so a second click while the retry is in-flight returns 409 instead of
    creating a duplicate task.
    """
    try:
        job = get_job(db, job_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="Job not found.")

    if job.status == JobStatus.PUBLISHING:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Publish already in progress.")

    if job.status != JobStatus.DONE:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Can only retry a completed job.")

    if not job.publish_failed:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Confluence publish did not fail for this job.")

    update_job_status(db, job_id, JobStatus.PUBLISHING, publish_failed=False)

    from tasks import retry_confluence_publish
    retry_confluence_publish.delay(job_id)

    return JSONResponse(
        status_code=status.HTTP_202_ACCEPTED,
        content={"job_id": job_id, "status": "retrying"},
    )


# ---------------------------------------------------------------------------
# GET /jobs/{job_id}/download
# ---------------------------------------------------------------------------

@app.get("/jobs/{job_id}/download", tags=["jobs"])
async def download_job_docx(
    job_id: str,
    api_key: str | None = None,
    db: Session = Depends(get_db),
) -> StreamingResponse:
    """
    Generate and stream a .docx for a completed job.

    The API key may be supplied as the ``api_key`` query param because this
    endpoint is opened via a browser anchor click (no header injection possible).
    """
    try:
        job = get_job(db, job_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="Job not found.")

    if job.status == JobStatus.FAILED:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Job failed.")

    if job.status != JobStatus.DONE:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Job is not complete yet.",
        )

    try:
        result = json.loads(job.result_json) if job.result_json else {}
    except json.JSONDecodeError:
        result = {}

    docx_bytes = _build_docx(result)

    raw_title = result.get("title", "meeting-notes")
    safe = re.sub(r"[^\w\s-]", "", raw_title).strip()
    safe = re.sub(r"\s+", "-", safe) or "meeting-notes"
    filename = f"{safe}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# .docx builder
# ---------------------------------------------------------------------------

def _build_docx(result: dict) -> bytes:
    """Build a Word document from the agent result dict and return raw bytes."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    output_type = result.get("output_type", "detailed")

    doc.add_heading(result.get("title") or "Meeting Notes", level=0)

    if output_type == "quick_summary":
        bullets = result.get("bullets") or []
        if bullets:
            doc.add_heading("Summary", level=1)
            for b in bullets:
                doc.add_paragraph(f"• {b}")
        action_items = result.get("action_items") or []
        if action_items:
            doc.add_heading("Action Items", level=1)
            for item in action_items:
                owner = item.get("owner") or "Unassigned"
                due = item.get("due") or "TBD"
                doc.add_paragraph(f"• {owner}: {item.get('task','')} (by {due})")

    elif output_type == "action_items":
        action_items = result.get("action_items") or []
        if action_items:
            doc.add_heading("Action Items", level=1)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Owner"
            hdr[1].text = "Task"
            hdr[2].text = "Due"
            for item in action_items:
                row = tbl.add_row().cells
                row[0].text = item.get("owner") or ""
                row[1].text = item.get("task") or ""
                row[2].text = item.get("due") or ""

    elif output_type == "mom":
        meeting_type = result.get("meeting_type", "")
        attendees = result.get("attendees") or []
        date = result.get("date", "")
        meta_parts = []
        if meeting_type:
            meta_parts.append(meeting_type.replace("-", " ").title())
        if date:
            meta_parts.append(f"Date: {date}")
        if attendees:
            meta_parts.append("Attendees: " + ", ".join(attendees))
        if meta_parts:
            p = doc.add_paragraph()
            run = p.add_run("  |  ".join(meta_parts))
            run.italic = True
            run.font.size = Pt(10)
        for heading, key in [("Agenda Items", "agenda_items"), ("Decisions", "decisions")]:
            items = result.get(key) or []
            if items:
                doc.add_heading(heading, level=1)
                for i, text in enumerate(items, 1):
                    doc.add_paragraph(f"{i}. {text}")
        action_items = result.get("action_items") or []
        if action_items:
            doc.add_heading("Action Items", level=1)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Owner"
            hdr[1].text = "Task"
            hdr[2].text = "Due"
            for item in action_items:
                row = tbl.add_row().cells
                row[0].text = item.get("owner") or ""
                row[1].text = item.get("task") or ""
                row[2].text = item.get("due") or ""
        next_steps = result.get("next_steps") or []
        if next_steps:
            doc.add_heading("Next Steps", level=1)
            for text in next_steps:
                doc.add_paragraph(f"• {text}")

    else:  # detailed
        meeting_type = result.get("meeting_type", "")
        attendees = result.get("attendees") or []
        meta_parts = []
        if meeting_type:
            meta_parts.append(meeting_type.replace("-", " ").title())
        if attendees:
            meta_parts.append("Attendees: " + ", ".join(attendees))
        if meta_parts:
            p = doc.add_paragraph()
            run = p.add_run("  |  ".join(meta_parts))
            run.italic = True
            run.font.size = Pt(10)
        if result.get("summary"):
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(result["summary"])
        decisions = result.get("decisions") or []
        if decisions:
            doc.add_heading("Decisions", level=1)
            for i, text in enumerate(decisions, 1):
                doc.add_paragraph(f"{i}. {text}")
        action_items = result.get("action_items") or []
        if action_items:
            doc.add_heading("Action Items", level=1)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Owner"
            hdr[1].text = "Task"
            hdr[2].text = "Due"
            for item in action_items:
                row = tbl.add_row().cells
                row[0].text = item.get("owner") or ""
                row[1].text = item.get("task") or ""
                row[2].text = item.get("due") or ""
        for heading, key in [("Open Questions", "open_questions"), ("Highlights", "highlights")]:
            items = result.get(key) or []
            if items:
                doc.add_heading(heading, level=1)
                for i, text in enumerate(items, 1):
                    doc.add_paragraph(f"{i}. {text}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port)
